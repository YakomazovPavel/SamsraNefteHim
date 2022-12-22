import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from progress.bar import IncrementalBar
import time


def cellStyle(cell):
    paragra_ph = cell.paragraphs[0]
    # Выравнивание по центру
    paragra_ph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    font = paragra_ph.runs[0].font
    # Шрифт
    font.name = 'Arial'
    # Размер
    font.size = Pt(9)


document = Document('Шаблон.docx')
table = document.tables[-1]


def addTableRow(array):
    row = table.add_row()
    cells = row.cells
    for index in range(0, len(cells)):
        cell = cells[index]
        cell.text = str(array[index])
        cellStyle(cell)


df = pd.ExcelFile('2022_12_06_База_МДЭА.xlsx').parse('База')
ColName = pd.read_excel('Имена столбцов.xlsx').head(1)
df.columns = ColName.columns
df = df.loc[3:]
df = df.fillna('').reset_index(drop=True)
df['№п/п'] = df['№п/п'].astype('int')
bar = IncrementalBar(f"Выполняется запись: ", max=df.shape[0])
for index1, row in df.iterrows():
    tableRow = list(row[['№п/п',
                         '№Блока',
                         'КодКонтура1',
                         'КодКонтура2',
                         'ПозицияКонтура',
                         'ПозицияПрибора',
                         'Hаименование',
                         'ТипСигнала',
                         'ТипСигнала2',
                         'БлокировкаLL',
                         'БлокировкаHH',
                         'СигнализацияLL',
                         'СигнализацияL',
                         'СигнализацияH',
                         'СигнализацияHH',
                         'ШкалаНиз',
                         'ШкалаВерх',
                         'ЕдиницаИзм',
                         'Клапан',
                         'Система',
                         'ШкафКроссовый',
                         'СхемаТехнологическая',
                         '№ЛистаТХ',
                         'Примечание'
                         ]].values)

    addTableRow(tableRow)
    bar.next()
# table.alignment = WD_TABLE_ALIGNMENT.CENTER
document.save('5766603-400-2-2-ATX-04-ПЭ-001.docx')
bar.finish()
print('Запись закончена!')
time.sleep(10)
